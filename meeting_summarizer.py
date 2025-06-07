import os
from datetime import datetime

import sounddevice as sd
import soundfile as sf
import openai
from docx import Document


def record_audio(duration: int, filename: str, samplerate: int = 44100, channels: int = 1) -> None:
    """Record audio from the microphone and save it to a WAV file."""
    print(f"Recording audio for {duration} seconds...")
    recording = sd.rec(int(duration * samplerate), samplerate=samplerate, channels=channels)
    sd.wait()
    sf.write(filename, recording, samplerate)
    print(f"Audio saved to {filename}")


def transcribe_audio(audio_file: str) -> str:
    """Transcribe the given audio file using OpenAI's Whisper API."""
    with open(audio_file, "rb") as f:
        response = openai.Audio.transcribe("whisper-1", f)
    return response["text"]


def extract_action_items(transcript: str, model: str = "gpt-3.5-turbo") -> str:
    """Use an LLM to extract action items and key decisions."""
    system_prompt = (
        "You are an assistant that summarizes meeting transcripts "
        "into action items and key decisions."
    )
    user_prompt = (
        "Extract the action items and key decisions from the following meeting transcript:\n" + transcript
    )

    response = openai.ChatCompletion.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0,
    )
    return response.choices[0].message["content"].strip()


def create_word_doc(transcript: str, summary: str, output_path: str) -> None:
    """Create a Word document containing the transcript and summary."""
    doc = Document()
    doc.add_heading("Meeting Minutes", level=1)

    doc.add_heading("Transcript", level=2)
    doc.add_paragraph(transcript)

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(summary)

    doc.save(output_path)
    print(f"Meeting minutes saved to {output_path}")


def main(duration: int, output_dir: str = "output") -> None:
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    audio_path = os.path.join(output_dir, f"meeting_{timestamp}.wav")
    record_audio(duration, audio_path)

    transcript = transcribe_audio(audio_path)
    summary = extract_action_items(transcript)

    doc_path = os.path.join(output_dir, f"meeting_minutes_{timestamp}.docx")
    create_word_doc(transcript, summary, doc_path)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="AI-Driven Meeting Minutes Summarizer")
    parser.add_argument(
        "--duration",
        type=int,
        default=60,
        help="Recording duration in seconds",
    )
    parser.add_argument(
        "--output-dir",
        default="output",
        help="Directory to store audio recordings and summaries",
    )

    args = parser.parse_args()
    main(args.duration, args.output_dir)
